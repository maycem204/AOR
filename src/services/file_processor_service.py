"""
Service de traitement des fichiers de connaissances pour l'indexation
"""

import logging
import os
from typing import List, Optional, Dict, Any
import re
from pathlib import Path
from .vector_store_service import VectorStoreService # ✅ Ajouter si pas déjà présent
from config.settings import settings
from models.data_models import Chunk, IndexingResult
from utils.file_utils import generate_chunk_id, is_supported_file, get_file_size_mb
from .excel_service import ExcelService

logger = logging.getLogger(__name__)
from .embedding_service import EmbeddingService  # ✅ Ajouter si pas déjà présent
class FileProcessorService:
    """
    Service pour le traitement des fichiers de connaissances
    """
    
    def __init__(self):
        """Initialise le service de traitement de fichiers"""
        self.excel_service = ExcelService()
        self.chunk_size = settings.CHUNK_SIZE
        self.embedding_service = EmbeddingService()  # ✅ Ajouter cette ligne
        self.vector_store_service = VectorStoreService()  # ✅ AJOUTER cette ligne
        self.chunk_overlap = settings.CHUNK_OVERLAP
        
        logger.info("Initialisation du service de traitement de fichiers")
    
    def process_knowledge_base(self) -> IndexingResult:
        """
        Traite tous les fichiers de la base de connaissances
        
        Returns:
            Résultat de l'indexation
        """
        start_time = time.time()
        files_processed = 0
        total_chunks = 0
        total_vectors = 0
        errors = []
        
        try:
            # Récupérer la liste des fichiers
            files = settings.get_knowledge_base_files()
            
            if not files:
                logger.warning("Aucun fichier trouvé dans la base de connaissances")
                return IndexingResult(
                    fichiers_traites=0,
                    chunks_crees=0,
                    vecteurs_stockes=0,
                    temps_indexation=time.time() - start_time,
                    erreurs=["Aucun fichier trouvé"]
                )
            
            logger.info(f"Traitement de {len(files)} fichiers de connaissances")
            
            for file_path in files:
                try:
                    logger.info(f"Traitement du fichier: {file_path}")
                    
                    # Extraire le texte du fichier
                    text_content = self._extract_text_from_file(file_path)
                    
                    if not text_content:
                        logger.warning(f"Aucun contenu extrait du fichier: {file_path}")
                        continue
                    
                    # Diviser en chunks
                    chunks = self._create_chunks(text_content, file_path)
                    
                    '''if chunks:
                        total_chunks += len(chunks)
                        files_processed += 1
                        logger.info(f"✓ {len(chunks)} chunks créés pour {file_path}")
                        '''
                    if chunks:
                    # Générer les embeddings
                        contents = [chunk.content for chunk in chunks]
                        embeddings = self.embedding_service.generate_embeddings_batch(contents)
    
                        for chunk, emb in zip(chunks, embeddings):
                            chunk.embedding = emb
    
                    # Filtrer les chunks valides
                        valid_chunks = [chunk for chunk in chunks if chunk.embedding is not None]
                        logger.info(f"✓ {len(valid_chunks)} / {len(chunks)} chunks avec embeddings valides")
    
                     # Insérer dans la base vectorielle
                        inserted = self.vector_store_service.insert_chunks(valid_chunks)
                        total_vectors += inserted
    
                        total_chunks += len(chunks)
                        files_processed += 1
                        logger.info(f"✓ {inserted} vecteurs insérés pour {file_path}")
                    #w houni rja3na lel code l 9dim
                    else:
                        logger.warning(f"Aucun chunk créé pour {file_path}")
                    
                except Exception as e:
                    error_msg = f"Erreur lors du traitement de {file_path}: {e}"
                    logger.error(error_msg)
                    errors.append(error_msg)
            
            temps_indexation = time.time() - start_time
            
            logger.info(f"✓ Indexation terminée: {files_processed} fichiers, {total_chunks} chunks")
            
            return IndexingResult(
                fichiers_traites=files_processed,
                chunks_crees=total_chunks,
                vecteurs_stockes=total_vectors,
                temps_indexation=temps_indexation,
                erreurs=errors
            )
            
        except Exception as e:
            error_msg = f"Erreur générale lors de l'indexation: {e}"
            logger.error(error_msg)
            return IndexingResult(
                fichiers_traites=files_processed,
                chunks_crees=total_chunks,
                vecteurs_stockes=total_vectors,
                temps_indexation=time.time() - start_time,
                erreurs=errors + [error_msg]
            )
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """
        Extrait le texte d'un fichier selon son type
        
        Args:
            file_path: Chemin du fichier
            
        Returns:
            Texte extrait du fichier
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension in settings.SUPPORTED_EXCEL_EXTENSIONS:
                return self.excel_service.extract_text_from_excel(file_path)
            
            elif file_extension in settings.SUPPORTED_WORD_EXTENSIONS:
                return self._extract_text_from_word(file_path)
            
            else:
                logger.warning(f"Type de fichier non supporté: {file_path}")
                return ""
                
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte de {file_path}: {e}")
            return ""
    
    def _extract_text_from_word(self, file_path: str) -> str:
        """
        Extrait le texte d'un fichier Word
        
        Args:
            file_path: Chemin du fichier Word
            
        Returns:
            Texte extrait du fichier Word
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extraire le texte des paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extraire le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            logger.info(f"✓ Texte extrait du fichier Word: {len(full_text)} caractères")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte Word: {e}")
            return ""
    
    def _create_chunks(self, text: str, source_file: str) -> List[Chunk]:
        """
        Divise un texte en chunks
        
        Args:
            text: Texte à diviser
            source_file: Fichier source
            
        Returns:
            Liste de chunks
        """
        try:
            if not text or len(text.strip()) < self.chunk_size:
                # Si le texte est plus petit que la taille de chunk, le traiter comme un seul chunk
                if text.strip():
                    chunk = Chunk(
                        id=generate_chunk_id(text, source_file, 0),
                        content=text.strip(),
                        source_file=source_file,
                        chunk_index=0,
                        embedding=None,  # Sera généré plus tard
                        metadata={
                            "file_size_mb": get_file_size_mb(source_file),
                            "chunk_type": "single"
                        }
                    )
                    return [chunk]
                return []
            
            chunks = []
            start = 0
            chunk_index = 0
            
            while start < len(text):
                # Définir la fin du chunk
                end = start + self.chunk_size
                
                # Si ce n'est pas le dernier chunk, essayer de couper à un espace
                if end < len(text):
                    # Chercher le dernier espace dans la fenêtre de chevauchement
                    overlap_start = max(start, end - self.chunk_overlap)
                    last_space = text.rfind(' ', overlap_start, end)
                    
                    if last_space > start:
                        end = last_space
                
                # Extraire le chunk
                chunk_text = text[start:end].strip()
                
                if chunk_text:
                    chunk = Chunk(
                        id=generate_chunk_id(chunk_text, source_file, chunk_index),
                        content=chunk_text,
                        source_file=source_file,
                        chunk_index=chunk_index,
                        embedding=None,  # Sera généré plus tard
                        metadata={
                            "file_size_mb": get_file_size_mb(source_file),
                            "chunk_type": "split",
                            "start_char": start,
                            "end_char": end,
                            "text_length": len(chunk_text)
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Passer au chunk suivant
                start = end - self.chunk_overlap
                
                # Éviter les boucles infinies
                if start >= len(text):
                    break
            
            logger.debug(f"✓ {len(chunks)} chunks créés pour {source_file}")
            return chunks
            
        except Exception as e:
            logger.error(f"Erreur lors de la création des chunks pour {source_file}: {e}")
            return []
    
    def process_single_file(self, file_path: str) -> List[Chunk]:
        """
        Traite un seul fichier et retourne ses chunks
        
        Args:
            file_path: Chemin du fichier à traiter
            
        Returns:
            Liste de chunks créés
        """
        try:
            logger.info(f"Traitement du fichier unique: {file_path}")
            
            # Vérifier que le fichier existe
            if not os.path.exists(file_path):
                logger.error(f"Fichier non trouvé: {file_path}")
                return []
            
            # Vérifier que le fichier est supporté
            if not is_supported_file(file_path, settings.SUPPORTED_EXTENSIONS):
                logger.error(f"Type de fichier non supporté: {file_path}")
                return []
            
            # Extraire le texte
            text_content = self._extract_text_from_file(file_path)
            
            if not text_content:
                logger.warning(f"Aucun contenu extrait du fichier: {file_path}")
                return []
            
            # Créer les chunks
            chunks = self._create_chunks(text_content, file_path)
            
            logger.info(f"✓ {len(chunks)} chunks créés pour {file_path}")
            return chunks
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du fichier {file_path}: {e}")
            return []
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """
        Récupère les statistiques de traitement
        
        Returns:
            Dictionnaire avec les statistiques
        """
        try:
            files = settings.get_knowledge_base_files()
            
            stats = {
                "total_files": len(files),
                "supported_files": len([f for f in files if is_supported_file(f, settings.SUPPORTED_EXTENSIONS)]),
                "excel_files": len([f for f in files if f.lower().endswith(('.xlsx', '.xls'))]),
                "word_files": len([f for f in files if f.lower().endswith(('.docx', '.doc'))]),
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap
            }
            
            return stats
            
        except Exception as e:
            logger.error(f"Erreur lors de la récupération des stats: {e}")
            return {"error": str(e)}

# Import time pour les calculs de temps
import time 